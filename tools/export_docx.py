"""
Convert docs/Project_Documentation.md (Markdown) to a .docx using python-docx.

Note: This is a simple converter that maps headings and paragraphs.
It does not support full Markdown (tables/code). For rich output,
consider pandoc. This script avoids extra OS-level dependencies.
"""

from pathlib import Path
from docx import Document


def md_to_docx(md_text: str) -> Document:
    doc = Document()
    lines = md_text.splitlines()

    def add_heading(text: str, level: int):
        level = max(1, min(level, 9))
        doc.add_heading(text.strip(), level=level)

    buffer: list[str] = []

    def flush_paragraph():
        if buffer:
            doc.add_paragraph("\n".join(buffer).strip())
            buffer.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            continue

        # Very simple heading detection based on leading hashes
        if line.startswith("#"):
            flush_paragraph()
            hashes = 0
            for ch in line:
                if ch == '#':
                    hashes += 1
                else:
                    break
            title = line[hashes:].strip(" #-")
            add_heading(title or " ", level=hashes)
        else:
            buffer.append(line)

    flush_paragraph()
    return doc


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md_path = root / "docs" / "Project_Documentation.md"
    out_path = root / "docs" / "Project_Documentation.docx"

    if not md_path.exists():
        raise SystemExit(f"Markdown not found: {md_path}")

    text = md_path.read_text(encoding="utf-8")
    document = md_to_docx(text)
    document.save(out_path)
    print(f"Saved: {out_path}")


